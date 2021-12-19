import PySimpleGUI as sg
from docx import Document
import openpyxl
from docx2pdf import convert
import os
from num2words import num2words

import win32gui, win32con



sg.theme('Light Blue 3')

progressbar = [
    [sg.ProgressBar(50, orientation='h', size=(51, 10), key='progress')]
]
outputwin = [
    [sg.Output(size=(78,20))]
]

layout = [[sg.Text('Invoice Generator')],
          [sg.Text('Select input Excel File:', size=(20, 1)), sg.Input(key='Input_1'), sg.FileBrowse()],
          [sg.Radio('PDF: ', "RADIO1", default=True, key="format")],
          [sg.Radio('Word: ', "RADIO1", default=False)],
          [sg.Frame('Progress',layout= progressbar)],
          [sg.Frame('Output', layout = outputwin)],
          [sg.Submit(), sg.Cancel()]]

window = sg.Window('Invoice Generator by Harshal', layout)
progress_bar = window.find_element('progress')
user_profile = os.environ['USERPROFILE']
user_desktop = user_profile + "\Documents\invoice_gen"
outputPath=os.getcwd()+'\\InvoiceGeneratedFiles'
inputPath=os.getcwd()+'\\ProgramFiles'

while True:
    event, values = window.read()
    if event in (sg.WINDOW_CLOSED, "Cancel"):
        break
    elif event == 'Submit':
        file_1 = values['Input_1']
        if '' not in [file_1]:
            print("hello")
            wb=openpyxl.load_workbook(file_1)

            progress_bar.UpdateBar(1, 5)

            print("reading file........")
            sh=wb['Sheet1']
            row=sh.max_row
            progress_bar.UpdateBar(2, 5)

            for i in range(2,row+1):
                invoice_date=str(sh.cell(i,1).value)
                recipient_name=sh.cell(i,2).value
                print("Reading data for "+recipient_name+".........")
                address=sh.cell(i,4).value
                invoice_number=sh.cell(i,6).value
                product_name=sh.cell(i,39).value
                GSTIN=sh.cell(i,8).value
                orderno_date=sh.cell(i,14).value
                mode=sh.cell(i,15).value
                other_ref=sh.cell(i,16).value
                term_delivery=sh.cell(i,17).value
                del_note_date=sh.cell(i,18).value
                dispatch_no=sh.cell(i,19).value
                dispatch_th=sh.cell(i,20).value
                destination=sh.cell(i,21).value
                gross_total=sh.cell(i,34).value
                word_total=num2words(gross_total)
                cgst=float(sh.cell(i,40).value)
                sgst=float(sh.cell(i,41).value)
                CGST_V=(gross_total*cgst)/100
                SGST_V=(gross_total*sgst)/100    
                t=sh.cell(i,11).value
                if(len(t)==0):
                    Pan_number="NA"
                else:
                    Pan_number=t   
                qty=sh.cell(i,30).value
                value=gross_total-(CGST_V+SGST_V)
                rate=value/qty

                print("Generating word formatted invoice......")
                document=Document(inputPath+'\\sample_invoice.docx')
                
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if "#Recipient_name" in run.text:
                                        run.text=run.text.replace('#Recipient_name',recipient_name)
                                    if "#Recipient_Address" in run.text:
                                        run.text=run.text.replace('#Recipient_Address',address)
                                    if "#Invoice_Number" in run.text:
                                        run.text=run.text.replace('#Invoice_Number',invoice_number)
                                    if "#GSTIN" in run.text:
                                        run.text=run.text.replace('#GSTIN',GSTIN)
                                    if "#Pan_Number" in run.text:
                                        run.text=run.text.replace('#Pan_Number',Pan_number)
                                    if "#Invoice_date" in run.text:
                                        run.text=run.text.replace('#Invoice_date',invoice_date)
                                    if "#Material_name" in run.text:
                                        run.text=run.text.replace('#Material_name',product_name)
                                    if "#orderno_date" in run.text:
                                        run.text=run.text.replace('#orderno_date',orderno_date)
                                    if "#mode" in run.text:
                                        run.text=run.text.replace('#mode',mode)
                                    if "#other_ref" in run.text:
                                        run.text=run.text.replace('#other_ref',other_ref)
                                    if "#term_delivery" in run.text:
                                        run.text=run.text.replace('#term_delivery',term_delivery)
                                    if "#del_note_date" in run.text:
                                        run.text=run.text.replace('#del_note_date',del_note_date)
                                    if "#dispatch_no" in run.text:
                                        run.text=run.text.replace('#dispatch_no',dispatch_no)
                                    if "#dispatch_th" in run.text:
                                        run.text=run.text.replace('#dispatch_th',dispatch_th)
                                    if "#destination" in run.text:
                                        run.text=run.text.replace('#destination',destination)
                                    if "#gross_total" in run.text:
                                        run.text=run.text.replace('#gross_total',str(gross_total))
                                    if "#CGST_V" in run.text:
                                        run.text=run.text.replace('#CGST_V',str(CGST_V))
                                    if "#SGST_V" in run.text:
                                        run.text=run.text.replace('#SGST_V',str(SGST_V))
                                    if "#value" in run.text:
                                        run.text=run.text.replace('#value',str(value))
                                    if "#rate" in run.text:
                                        run.text=run.text.replace('#rate',str(rate))
                                    if "#qty" in run.text:
                                        run.text=run.text.replace('#qty',str(qty))
                                    if "#cgst" in run.text:
                                        run.text=run.text.replace('#cgst',str(cgst))
                                    if "#sgst" in run.text:
                                        run.text=run.text.replace('#sgst',str(sgst))
                                    if "#word_total" in run.text:
                                        run.text=run.text.replace('#word_total',str(word_total))

                            fileN=recipient_name+invoice_number[-1]
                            
                            document.save(outputPath+"\\"+fileN+'.docx')

            progress_bar.UpdateBar(3, 5)

            print("Generating pdf formatted invoices. Please wait......")

            print(values["format"]);

            convert(outputPath+'\\')
        
            progress_bar.UpdateBar(4, 5)
            if(values["format"]==True):
                print("Deleting Word formatted invoices...")
                directory = outputPath+"\\"
                files_in_directory = os.listdir(directory)
                filtered_files = [file for file in files_in_directory if file.endswith(".docx")]
                for file in filtered_files:
                    path_to_file = os.path.join(directory, file)
                    os.remove(path_to_file)
            print("PDF formatted invoices created successfully")
            print("Please find invoices in path: "+outputPath)
            progress_bar.UpdateBar(5, 5)
            
window.close()
